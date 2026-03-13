from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime


def generar_documento(titulo: str, contenido: str) -> BytesIO:
    """
    Genera un documento .docx con el título y contenido recibidos.
    Retorna un objeto BytesIO (el archivo en memoria, sin guardarlo en disco).
    """

    doc = Document()

    # ── Márgenes ──────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Fecha (pequeña, arriba a la derecha) ──────────────────
    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fecha_run = fecha_p.add_run(datetime.now().strftime("%d/%m/%Y"))
    fecha_run.font.size = Pt(9)
    fecha_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── Título ────────────────────────────────────────────────
    titulo_p = doc.add_heading(level=1)
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo_p.add_run(titulo)
    titulo_run.font.size = Pt(22)
    titulo_run.font.bold = True
    titulo_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # azul oscuro

    # Línea separadora
    doc.add_paragraph("─" * 60)

    # ── Contenido ─────────────────────────────────────────────
    # Dividimos por saltos de línea para respetar párrafos
    parrafos = contenido.strip().split("\n")
    for parrafo in parrafos:
        parrafo = parrafo.strip()
        if not parrafo:
            doc.add_paragraph("")  # línea en blanco
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(parrafo)
        run.font.size = Pt(12)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(8)

    # ── Pie de página ─────────────────────────────────────────
    doc.add_paragraph("")
    pie_p = doc.add_paragraph()
    pie_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie_run = pie_p.add_run("Documento generado automáticamente")
    pie_run.font.size = Pt(8)
    pie_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    pie_run.font.italic = True

    # ── Guardar en memoria ────────────────────────────────────
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
